import os
from datetime import datetime
from docx import Document
from docx.shared import Cm
from PIL import Image
import streamlit as st
import tempfile
import time
import io

def reduzir_imagem(imagem_bytes, largura_cm, altura_cm):
    with Image.open(imagem_bytes) as img:
        dpi = img.info.get("dpi", (96, 96))[0]
        largura_px = int((largura_cm / 2.54) * dpi)
        altura_px = int((altura_cm / 2.54) * dpi)
        img = img.resize((largura_px, altura_px), Image.LANCZOS)
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
        img.save(temp_file.name)
        return temp_file.name

def criar_interface_mobile_friendly():
    """Cria uma interface otimizada para mobile"""
    st.markdown("""
    <style>
    /* Otimizações para mobile */
    .stFileUploader > div > div > div > div {
        padding: 1rem !important;
    }
    .mobile-warning {
        background-color: #fff3e0;
        padding: 15px;
        border-radius: 10px;
        border-left: 4px solid #ff9800;
        margin: 10px 0;
        font-size: 14px;
    }
    .success-box {
        background-color: #e8f5e8;
        padding: 15px;
        border-radius: 10px;
        border-left: 4px solid #4caf50;
        margin: 10px 0;
    }
    .info-box {
        background-color: #e3f2fd;
        padding: 15px;
        border-radius: 10px;
        border-left: 4px solid #2196f3;
        margin: 10px 0;
    }
    /* Melhor visualização em telas pequenas */
    @media (max-width: 768px) {
        .stButton > button {
            width: 100% !important;
            font-size: 16px !important;
            padding: 12px !important;
        }
        .stFileUploader {
            margin-bottom: 20px !important;
        }
    }
    </style>
    """, unsafe_allow_html=True)

def salvar_fotos_session_state(fotos, chave):
    """Salva as fotos no session state para persistência"""
    if fotos:
        fotos_data = []
        for foto in fotos:
            foto_data = {
                'name': foto.name,
                'size': foto.size,
                'type': foto.type,
                'data': foto.getvalue()
            }
            fotos_data.append(foto_data)
        st.session_state[chave] = fotos_data
        return True
    return False

def recuperar_fotos_session_state(chave):
    """Recupera as fotos do session state"""
    return st.session_state.get(chave, [])

def inserir_bloco_imagens(doc, titulo, imagens_data, largura_cm=5, altura_cm=4):
    """Insere bloco de imagens no documento. Aceita tanto arquivos quanto dados do session state"""
    doc.add_paragraph("------------------------------------------")
    doc.add_heading(titulo, level=2)
    par = doc.add_paragraph()
    
    for imagem in imagens_data:
        if isinstance(imagem, dict):  # Dados do session state
            imagem_bytes = io.BytesIO(imagem['data'])
        else:  # UploadedFile normal
            imagem_bytes = imagem
            
        img_path = reduzir_imagem(imagem_bytes, largura_cm, altura_cm)
        par.add_run().add_picture(img_path, width=Cm(largura_cm), height=Cm(altura_cm))
        os.remove(img_path)

# Configuração da página
st.set_page_config(
    page_title="Relatório Zeladoria",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS e interface mobile-friendly
criar_interface_mobile_friendly()

# Inicializar session state
if 'site_id' not in st.session_state:
    st.session_state.site_id = ""
if 'data_execucao' not in st.session_state:
    st.session_state.data_execucao = datetime.now().date()
if 'localizacao' not in st.session_state:
    st.session_state.localizacao = ""

st.title("📋 Relatório Fotográfico de Zeladoria")

# Aviso importante para mobile
st.markdown("""
<div class="mobile-warning">
    <strong>📱 IMPORTANTE - Usuários de celular:</strong><br>
    • Preencha TODOS os campos de texto primeiro<br>
    • Faça upload de UMA categoria de fotos por vez<br>
    • Aguarde a confirmação "✅ carregada com sucesso" antes do próximo upload<br>
    • NÃO saia do navegador durante o upload<br>
    • Se der "reconnecting", recarregue a página e tente novamente
</div>
""", unsafe_allow_html=True)

# Formulário com campos persistentes
st.subheader("📝 Informações do Relatório")

col1, col2 = st.columns(2)
with col1:
    site_id = st.text_input(
        "ID do site", 
        value=st.session_state.site_id,
        key="input_site_id",
        help="Digite o ID do site"
    )
    if site_id != st.session_state.site_id:
        st.session_state.site_id = site_id

with col2:
    data_execucao = st.date_input(
        "Data da execução", 
        value=st.session_state.data_execucao,
        key="input_data"
    )
    if data_execucao != st.session_state.data_execucao:
        st.session_state.data_execucao = data_execucao

localizacao = st.text_input(
    "Localização (cidade - estado)", 
    value=st.session_state.localizacao,
    key="input_localizacao",
    help="Digite a localização completa"
)
if localizacao != st.session_state.localizacao:
    st.session_state.localizacao = localizacao

# Status dos dados salvos
if st.session_state.site_id and st.session_state.localizacao:
    st.success("✅ Dados do formulário salvos!")

st.divider()

# Upload de fotos com persistência
st.subheader("📸 Upload de Fotos")

# Instruções específicas para mobile
st.markdown("""
<div class="info-box">
    <strong>💡 Dica:</strong> No celular, toque em "Browse files" → Escolha "Câmera" ou "Galeria" → Selecione as fotos → Aguarde o upload completar
</div>
""", unsafe_allow_html=True)

# Seção ANTES
st.markdown("### 🔴 Fotos ANTES")
fotos_antes = st.file_uploader(
    "📸 Selecione as fotos do ANTES", 
    type=["jpg", "jpeg", "png"], 
    accept_multiple_files=True,
    key="upload_antes",
    help="Você pode selecionar múltiplas fotos de uma vez"
)

if fotos_antes:
    if salvar_fotos_session_state(fotos_antes, 'fotos_antes_data'):
        st.markdown(f"<div class='success-box'>✅ {len(fotos_antes)} foto(s) ANTES carregada(s) com sucesso!</div>", unsafe_allow_html=True)
        
        # Preview das fotos
        if len(fotos_antes) <= 4:
            cols = st.columns(min(len(fotos_antes), 4))
            for i, foto in enumerate(fotos_antes):
                with cols[i]:
                    st.image(foto, caption=f"Antes {i+1}", width=150)

# Verificar fotos salvas na sessão
fotos_antes_salvas = recuperar_fotos_session_state('fotos_antes_data')
if fotos_antes_salvas and not fotos_antes:
    st.markdown(f"<div class='info-box'>📁 {len(fotos_antes_salvas)} foto(s) ANTES já salvas na sessão</div>", unsafe_allow_html=True)

st.divider()

# Seção DEPOIS
st.markdown("### 🟢 Fotos DEPOIS")
fotos_depois = st.file_uploader(
    "📸 Selecione as fotos do DEPOIS", 
    type=["jpg", "jpeg", "png"], 
    accept_multiple_files=True,
    key="upload_depois",
    help="Você pode selecionar múltiplas fotos de uma vez"
)

if fotos_depois:
    if salvar_fotos_session_state(fotos_depois, 'fotos_depois_data'):
        st.markdown(f"<div class='success-box'>✅ {len(fotos_depois)} foto(s) DEPOIS carregada(s) com sucesso!</div>", unsafe_allow_html=True)
        
        # Preview das fotos
        if len(fotos_depois) <= 4:
            cols = st.columns(min(len(fotos_depois), 4))
            for i, foto in enumerate(fotos_depois):
                with cols[i]:
                    st.image(foto, caption=f"Depois {i+1}", width=150)

# Verificar fotos salvas na sessão
fotos_depois_salvas = recuperar_fotos_session_state('fotos_depois_data')
if fotos_depois_salvas and not fotos_depois:
    st.markdown(f"<div class='info-box'>📁 {len(fotos_depois_salvas)} foto(s) DEPOIS já salvas na sessão</div>", unsafe_allow_html=True)

st.divider()

# Seção PLACA
st.markdown("### 🏷️ Foto da Placa")
foto_placa = st.file_uploader(
    "📸 Selecione a foto da PLACA DE IDENTIFICAÇÃO", 
    type=["jpg", "jpeg", "png"],
    key="upload_placa",
    help="Apenas uma foto da placa"
)

if foto_placa:
    if salvar_fotos_session_state([foto_placa], 'foto_placa_data'):
        st.markdown("<div class='success-box'>✅ Foto da PLACA carregada com sucesso!</div>", unsafe_allow_html=True)
        st.image(foto_placa, caption="Placa de Identificação", width=200)

# Verificar foto da placa salva na sessão
foto_placa_salva = recuperar_fotos_session_state('foto_placa_data')
if foto_placa_salva and not foto_placa:
    st.markdown("<div class='info-box'>📁 Foto da PLACA já salva na sessão</div>", unsafe_allow_html=True)

st.divider()

# Botão de gerar relatório
st.subheader("📄 Gerar Relatório")

# Verificar se tem dados suficientes
tem_dados_basicos = st.session_state.site_id and st.session_state.localizacao
tem_fotos = (fotos_antes or fotos_antes_salvas or 
             fotos_depois or fotos_depois_salvas or 
             foto_placa or foto_placa_salva)

if not tem_dados_basicos:
    st.warning("⚠️ Preencha os campos Site ID e Localização primeiro")
elif not tem_fotos:
    st.warning("⚠️ Faça upload de pelo menos uma foto")
else:
    col1, col2 = st.columns([3, 1])
    with col1:
        if st.button("🚀 Gerar Relatório", type="primary", use_container_width=True):
            try:
                with st.spinner("Gerando relatório..."):
                    # Usar dados do session state se as fotos não estiverem carregadas no momento
                    fotos_antes_final = fotos_antes if fotos_antes else fotos_antes_salvas
                    fotos_depois_final = fotos_depois if fotos_depois else fotos_depois_salvas
                    foto_placa_final = foto_placa if foto_placa else (foto_placa_salva[0] if foto_placa_salva else None)
                    
                    doc = Document()
                    doc.add_heading("RELATÓRIO FOTOGRÁFICO DE ZELADORIA", level=1)
                    doc.add_paragraph(f"Site ID: {st.session_state.site_id}")
                    doc.add_paragraph(f"Data da Execução: {st.session_state.data_execucao.strftime('%d/%m/%Y')}")
                    doc.add_paragraph(f"Localização: {st.session_state.localizacao.upper()}")

                    if fotos_antes_final:
                        inserir_bloco_imagens(doc, "FOTOS - ANTES", fotos_antes_final)
                    if fotos_depois_final:
                        inserir_bloco_imagens(doc, "FOTOS - DEPOIS", fotos_depois_final)
                    if foto_placa_final:
                        inserir_bloco_imagens(doc, "PLACA DE IDENTIFICAÇÃO", [foto_placa_final])

                    nome_arquivo = f"RLT. ZELADORIA - {st.session_state.site_id} - {st.session_state.data_execucao.strftime('%Y-%m-%d')}.docx"
                    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                    doc.save(temp_docx.name)

                    with open(temp_docx.name, "rb") as file:
                        st.success("✅ Relatório gerado com sucesso!")
                        st.download_button(
                            "📥 Baixar Relatório", 
                            file.read(), 
                            file_name=nome_arquivo,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary",
                            use_container_width=True
                        )
                    
                    # Limpar arquivos temporários
                    os.unlink(temp_docx.name)
                        
            except Exception as e:
                st.error(f"❌ Erro ao gerar relatório: {str(e)}")
                st.info("💡 Tente recarregar a página e fazer upload das fotos novamente")

    with col2:
        if st.button("🗑️ Limpar", help="Limpar todos os dados"):
            # Limpar session state
            for key in ['site_id', 'localizacao', 'fotos_antes_data', 'fotos_depois_data', 'foto_placa_data']:
                if key in st.session_state:
                    del st.session_state[key]
            st.session_state.site_id = ""
            st.session_state.localizacao = ""
            st.rerun()

# Sidebar com status
with st.sidebar:
    st.header("📊 Status da Sessão")
    
    # Indicador de conexão
    st.markdown("🟢 **Aplicação Online**")
    st.markdown(f"🕐 **{datetime.now().strftime('%H:%M:%S')}**")
    
    st.divider()
    
    st.subheader("📝 Dados do Formulário")
    st.write(f"🆔 Site ID: {'✅' if st.session_state.site_id else '❌'}")
    st.write(f"📍 Localização: {'✅' if st.session_state.localizacao else '❌'}")
    st.write(f"📅 Data: {st.session_state.data_execucao.strftime('%d/%m/%Y')}")
    
    st.subheader("📸 Fotos na Sessão")
    fotos_antes_count = len(fotos_antes_salvas) if fotos_antes_salvas else 0
    fotos_depois_count = len(fotos_depois_salvas) if fotos_depois_salvas else 0
    foto_placa_count = len(foto_placa_salva) if foto_placa_salva else 0
    
    st.write(f"🔴 Antes: {fotos_antes_count}")
    st.write(f"🟢 Depois: {fotos_depois_count}")
    st.write(f"🏷️ Placa: {foto_placa_count}")
    
    total_fotos = fotos_antes_count + fotos_depois_count + foto_placa_count
    st.metric("📊 Total", total_fotos)
    
    if tem_dados_basicos and total_fotos > 0:
        st.success("✅ Pronto!")
    elif not tem_dados_basicos:
        st.warning("⚠️ Faltam dados")
    else:
        st.info("ℹ️ Aguardando fotos")
        
    st.divider()
    
    # Dicas para mobile
    with st.expander("📱 Dicas para Mobile"):
        st.markdown("""
        **Se der erro "reconnecting":**
        1. Recarregue a página (F5)
        2. Preencha os dados novamente
        3. Faça upload UMA categoria por vez
        4. Aguarde confirmação antes da próxima
        5. Gere o relatório imediatamente
        
        **Para melhor experiência:**
        - Use WiFi ao invés de dados móveis
        - Feche outros apps durante o upload
        - Mantenha a tela ligada
        """)