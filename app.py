import os
from datetime import datetime
from docx import Document
from docx.shared import Cm
from PIL import Image
import streamlit as st
import tempfile
import time

def reduzir_imagem(imagem_bytes, largura_cm, altura_cm):
    with Image.open(imagem_bytes) as img:
        dpi = img.info.get("dpi", (96, 96))[0]
        largura_px = int((largura_cm / 2.54) * dpi)
        altura_px = int((altura_cm / 2.54) * dpi)
        img = img.resize((largura_px, altura_px), Image.LANCZOS)
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
        img.save(temp_file.name)
        return temp_file.name

def inserir_bloco_imagens(doc, titulo, imagens, largura_cm=5, altura_cm=4):
    doc.add_paragraph("------------------------------------------")
    doc.add_heading(titulo, level=2)
    par = doc.add_paragraph()
    for imagem in imagens:
        img_path = reduzir_imagem(imagem, largura_cm, altura_cm)
        par.add_run().add_picture(img_path, width=Cm(largura_cm), height=Cm(altura_cm))
        os.remove(img_path)

# Configurações da página
st.set_page_config(
    page_title="Relatório Fotográfico de Zeladoria",
    page_icon="📋",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# Inicializar session state para manter dados
if 'form_data' not in st.session_state:
    st.session_state.form_data = {
        'site_id': '',
        'data_execucao': datetime.now().date(),
        'localizacao': '',
        'fotos_antes': [],
        'fotos_depois': [],
        'foto_placa': None
    }

# Função para salvar dados no session state
def salvar_dados():
    st.session_state.form_data.update({
        'site_id': st.session_state.get('site_id_input', ''),
        'localizacao': st.session_state.get('localizacao_input', ''),
        'data_execucao': st.session_state.get('data_input', datetime.now().date())
    })

st.title("📋 Relatório Fotográfico de Zeladoria")
st.markdown("---")

# Instruções para o usuário
with st.expander("ℹ️ Instruções Importantes", expanded=False):
    st.markdown("""
    **Para evitar perda de dados:**
    1. Preencha os campos de texto primeiro
    2. Faça upload das fotos uma categoria por vez
    3. Gere o relatório logo após fazer todos os uploads
    4. Se a página recarregar, seus uploads serão perdidos
    """)

# Status da conexão
st.sidebar.markdown("🟢 **Aplicação Online**")
st.sidebar.markdown(f"⏰ **Última atualização:** {datetime.now().strftime('%H:%M:%S')}")

# Formulário principal com campos persistentes
col1, col2 = st.columns(2)

with col1:
    site_id = st.text_input(
        "ID do site",
        value=st.session_state.form_data['site_id'],
        key='site_id_input',
        on_change=salvar_dados,
        help="Digite o identificador único do site"
    )

with col2:
    data_execucao = st.date_input(
        "Data da execução",
        value=st.session_state.form_data['data_execucao'],
        key='data_input',
        on_change=salvar_dados
    )

localizacao = st.text_input(
    "Localização (cidade - estado)",
    value=st.session_state.form_data['localizacao'],
    key='localizacao_input',
    on_change=salvar_dados,
    help="Ex: São Paulo - SP"
)

st.markdown("---")
st.subheader("📸 Upload de Fotos")

# Upload de fotos com validação e feedback
col_antes, col_depois, col_placa = st.columns(3)

with col_antes:
    st.markdown("**ANTES**")
    fotos_antes = st.file_uploader(
        "Fotos do ANTES",
        type=["jpg", "jpeg", "png"],
        accept_multiple_files=True,
        key="fotos_antes_uploader",
        help="Selecione múltiplas fotos do estado anterior"
    )
    if fotos_antes:
        st.success(f"✅ {len(fotos_antes)} foto(s) carregada(s)")
        for i, foto in enumerate(fotos_antes[:3]):  # Mostrar preview das primeiras 3
            st.image(foto, width=100, caption=f"Antes {i+1}")

with col_depois:
    st.markdown("**DEPOIS**")
    fotos_depois = st.file_uploader(
        "Fotos do DEPOIS",
        type=["jpg", "jpeg", "png"],
        accept_multiple_files=True,
        key="fotos_depois_uploader",
        help="Selecione múltiplas fotos do estado posterior"
    )
    if fotos_depois:
        st.success(f"✅ {len(fotos_depois)} foto(s) carregada(s)")
        for i, foto in enumerate(fotos_depois[:3]):  # Mostrar preview das primeiras 3
            st.image(foto, width=100, caption=f"Depois {i+1}")

with col_placa:
    st.markdown("**PLACA DE ID**")
    foto_placa = st.file_uploader(
        "Foto da PLACA",
        type=["jpg", "jpeg", "png"],
        key="foto_placa_uploader",
        help="Foto da placa de identificação do local"
    )
    if foto_placa:
        st.success("✅ Placa carregada")
        st.image(foto_placa, width=100, caption="Placa")

st.markdown("---")

# Validação antes de gerar relatório
pode_gerar = bool(site_id and localizacao and (fotos_antes or fotos_depois or foto_placa))

if not pode_gerar:
    st.warning("⚠️ Preencha pelo menos: ID do site, localização e uma foto para gerar o relatório.")

# Botão de geração com validação
if st.button(
    "🔄 Gerar Relatório",
    disabled=not pode_gerar,
    help="Clique para gerar o relatório em formato Word" if pode_gerar else "Preencha os campos obrigatórios primeiro"
):
    try:
        with st.spinner("📄 Gerando relatório..."):
            # Salvar dados atuais no session state
            st.session_state.form_data.update({
                'site_id': site_id,
                'data_execucao': data_execucao,
                'localizacao': localizacao,
                'fotos_antes': fotos_antes,
                'fotos_depois': fotos_depois,
                'foto_placa': foto_placa
            })
            
            doc = Document()
            doc.add_heading("RELATÓRIO FOTOGRÁFICO DE ZELADORIA", level=1)
            doc.add_paragraph(f"Site ID: {site_id}")
            doc.add_paragraph(f"Data da Execução: {data_execucao.strftime('%d/%m/%Y')}")
            doc.add_paragraph(f"Localização: {localizacao.upper()}")

            if fotos_antes:
                inserir_bloco_imagens(doc, "FOTOS - ANTES", fotos_antes)
            if fotos_depois:
                inserir_bloco_imagens(doc, "FOTOS - DEPOIS", fotos_depois)
            if foto_placa:
                inserir_bloco_imagens(doc, "PLACA DE IDENTIFICAÇÃO", [foto_placa])

            nome_arquivo = f"RLT. ZELADORIA - {site_id} - {data_execucao.strftime('%Y-%m-%d')}.docx"
            temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            doc.save(temp_docx.name)

            with open(temp_docx.name, "rb") as file:
                st.success("✅ Relatório gerado com sucesso!")
                st.download_button(
                    "📥 Baixar Relatório",
                    file,
                    file_name=nome_arquivo,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            # Limpar arquivo temporário
            os.remove(temp_docx.name)
            
    except Exception as e:
        st.error(f"❌ Erro ao gerar relatório: {str(e)}")
        st.info("💡 Tente novamente ou recarregue a página se o problema persistir.")

# Rodapé com informações úteis
st.markdown("---")
st.markdown(
    "<div style='text-align: center; color: gray; font-size: 0.8em;'>"
    "💡 Mantenha esta aba aberta durante o upload das fotos para evitar perda de dados"
    "</div>",
    unsafe_allow_html=True
)
